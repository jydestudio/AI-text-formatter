import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import google.generativeai as genai
import time, io

st.set_page_config(
    page_title="AI text formatter",
)

st.title("My AI Text Formatter (Na soft life I wan dey live 😋)")

user_name = st.text_input("Full Name:", key="user_name")
user_matric = st.text_input("Matric Number", key="user_matric")
user_department = st.text_input("Department:", key="user_department")
user_school = st.text_input("School:", key="user_school")
user_course = st.text_input("Course:", key="user_course")

pages_options = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10]
pages = st.select_slider("Number of pages", options=pages_options, key="pages")

heading_font_options = [
    "Helvetica",
    "Arial",
    "Montserrat",
    "Lato",
    "Roboto",
    "Open Sans",
    "Poppins",
    "Source Sans Pro",
    "Futura",
    "Avenir",
    "Proxima Nova",
    "Raleway",
    "Nunito",
    "Quicksand",
    "Inter"
]

body_font_options = [
    "Electra",
    "Minion",
    "Sabon",
    "Calluna",
    "Neue Haas Unica",
    "Tiempos Text",
    "Freight Text",
    "Plantin",
    "Baskerville",
    "Acumin"
]

heading_font = st.selectbox("Select heading font", options=heading_font_options, key="heading_font")
body_font = st.selectbox("Select body font", options=body_font_options, key="body_font")

questions = st.text_area("Paste your questions here", key="questions")
make_standard = st.checkbox("Add intro and conclusion", value=False)    
use_citation = st.checkbox("Use citation", value=False)

st.markdown(
    """
    <style>
    .stDownloadButton>button {
        background-color: rgb(34, 139, 34);
        color: white;
    }

    .stDownloadButton>button:hover {
        background-color: transparent;
        color: white;
    }
    .stButton>button:hover {
        background-color: transparent;
        color: white;
    }

    .stButton>button {
        background-color: rgb(1, 153, 255);
        color: white;
    }

    </style>
    """, 
unsafe_allow_html=True)

def set_font(run, font_name, size, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold

doc = Document()
all_content = ""
if st.button("Generate DOCX"):
    placeholder = st.empty()
    placeholder.info("Generating document...")
    name = doc.add_paragraph()
    name_run = name.add_run(f'NAME: {user_name.upper()}\nMATRIC: {user_matric.upper()}')
    set_font(name_run, heading_font, 16, bold=True)
    name.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    dept = doc.add_paragraph()
    if len(user_school) > 6:
        user_school = user_school.title()
    else:
        user_school = user_school.upper()

    if len(user_department) > 6:
        user_department = user_department.title()
    else:
        user_department = user_department.upper()

    dept_run = dept.add_run(f"Department: {user_department}\nSchool: {user_school}\nCourse: {user_course.upper()}")
    set_font(dept_run, heading_font, 14)
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph()
    doc.add_paragraph()

    api_key = "AIzaSyCZP8cNH0ZA4zxYRp237UgkiOcxkQzSf4c"
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash")

    response = model.generate_content(f"""
            You only goal is determine the questions in a paragraph 
            and return the questions seperated using '###' here is the paragraph: 
            '{questions}'. your response
            should be the questions in this: '{questions}' seperated with '###'
            and nothing else.
            the content provided might not come with questions mark, but assume that 
            every thing in the content contain at least one question, determine how
            many questions are there and return the question in it with each separated
            using '###' your response should not contain any other thing apart from 
            each questions separeted using '###'.
            the question you should return should just get the topic from the question 
            and should not add words like 'define' or 'describe' or questions marks
            the question you should return should be like a heading name that answers
            the question
            """)

    questions = response.text.split("###")
    no_of_words_per_question = ((pages * 450) / len(questions))
    if use_citation:
        citation_command = "Write the content with valid inline citations using APA 7th edition format. Ensure the citations follow the structure (Author, Year). Use realistic and relevant references. note: do not quote the authors, only draw meaningful conclusion from their work and reference them appropriately"
    else:
        citation_command = ""

    if make_standard:
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"Introduction")
        set_font(subtitle_run, heading_font, 14, bold=True)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Content section
        content = doc.add_paragraph()
        response = model.generate_content(f"""
            i you to write a STRICTLY {no_of_words_per_question} words introduction on these topic in a context that match them together, here are the topics {",".join(questions)}, 
            the topics are subheading in the paper and i want you to create an introduction for it.
            you response should be the introduction and nothing else.
            
        """)

        content_run = content.add_run(response.text)

        set_font(content_run, body_font, 12)
        content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph()
    for index, question in enumerate(questions):
        # subtitle 
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"{question.title()}")
        set_font(subtitle_run, heading_font, 14, bold=True)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Content section
        content = doc.add_paragraph()
        response = model.generate_content(f"""
            i you to write a STRICTLY {no_of_words_per_question} words on {question}, 
            i want you to make it clear, use keywords and make it easy to understand.
            you response should be only content for the question, nothing else.
            Note: {citation_command}
        """)
        all_content += f"{response.text}\n\n"
        content_run = content.add_run(response.text)

        set_font(content_run, body_font, 12)
        content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        if len(questions) >= 15:
            time.sleep(4.5)

    doc.add_paragraph()
    doc.add_paragraph()

    if make_standard:
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"Conclusion")
        set_font(subtitle_run, heading_font, 14, bold=True)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Content section
        content = doc.add_paragraph()
        response = model.generate_content(f"""
            i want you to write a STRICTLY {no_of_words_per_question} words conclusion for these topics in a context that match them together, here are the topics {",".join(questions)}, 
            with the content {all_content}
            
        """)

        content_run = content.add_run(response.text)

        set_font(content_run, body_font, 12)
        content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


    for index, question in enumerate(questions):
        # subtitle 
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"{question.title()}")
        set_font(subtitle_run, heading_font, 14, bold=True)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc_data = io.BytesIO()
    doc.save(doc_data)


    placeholder.success("Document Generated, click the button below to download!")

    st.download_button(
        label="↓ Download DOCX file",
        data=doc_data.getvalue(),
        file_name="assignment.docx",
        mime="docx"
)
